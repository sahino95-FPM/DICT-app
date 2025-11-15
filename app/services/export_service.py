"""Service d'export de données"""
import csv
import io
from datetime import datetime
from flask import make_response
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class ExportService:
    """Service pour l'export de données dans différents formats"""

    @staticmethod
    def export_to_csv(data, columns, filename_prefix='export'):
        """
        Exporte les données en CSV

        Args:
            data: Liste de dictionnaires
            columns: Liste des colonnes à exporter
            filename_prefix: Préfixe du nom de fichier

        Returns:
            Response Flask avec le fichier CSV
        """
        timestamp = datetime.now().strftime('%Y-%m-%d_%H%M')
        filename = f"{filename_prefix}_{timestamp}.csv"

        # Création du CSV en mémoire
        si = io.StringIO()
        writer = csv.DictWriter(si, fieldnames=columns, extrasaction='ignore')

        writer.writeheader()
        writer.writerows(data)

        # Création de la réponse
        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename={filename}"
        output.headers["Content-type"] = "text/csv; charset=utf-8"

        return output

    @staticmethod
    def export_to_xlsx(data, columns, column_labels, filename_prefix='export', sheet_name='Données'):
        """
        Exporte les données en Excel (XLSX)

        Args:
            data: Liste de dictionnaires
            columns: Liste des clés des colonnes
            column_labels: Dict mapping clés -> libellés
            filename_prefix: Préfixe du nom de fichier
            sheet_name: Nom de la feuille

        Returns:
            Response Flask avec le fichier XLSX
        """
        timestamp = datetime.now().strftime('%Y-%m-%d_%H%M')
        filename = f"{filename_prefix}_{timestamp}.xlsx"

        # Création du workbook
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Style de l'en-tête
        header_fill = PatternFill(start_color="006b01", end_color="006b01", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        header_alignment = Alignment(horizontal="center", vertical="center")

        # Écriture de l'en-tête
        for col_idx, col_key in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx)
            cell.value = column_labels.get(col_key, col_key)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment

        # Écriture des données
        for row_idx, row_data in enumerate(data, start=2):
            for col_idx, col_key in enumerate(columns, start=1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = row_data.get(col_key)

                # Alignement selon le type
                if isinstance(cell.value, (int, float)):
                    cell.alignment = Alignment(horizontal="right")

        # Ajustement automatique de la largeur des colonnes
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column].width = adjusted_width

        # Sauvegarde en mémoire
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        # Création de la réponse
        response = make_response(output.read())
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        return response

    @staticmethod
    def export_to_pdf(data, columns, column_labels, title, filename_prefix='export',
                     metadata=None, logo_path=None):
        """
        Exporte les données en PDF

        Args:
            data: Liste de dictionnaires
            columns: Liste des clés des colonnes
            column_labels: Dict mapping clés -> libellés
            title: Titre du document
            filename_prefix: Préfixe du nom de fichier
            metadata: Dict avec métadonnées (période, filtres, etc.)
            logo_path: Chemin vers le logo

        Returns:
            Response Flask avec le fichier PDF
        """
        timestamp = datetime.now().strftime('%Y-%m-%d_%H%M')
        filename = f"{filename_prefix}_{timestamp}.pdf"

        # Création du PDF en mémoire
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            rightMargin=1*cm,
            leftMargin=1*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        elements = []
        styles = getSampleStyleSheet()

        # Logo (si disponible)
        if logo_path and os.path.exists(logo_path):
            try:
                logo = Image(logo_path, width=3*cm, height=3*cm)
                elements.append(logo)
                elements.append(Spacer(1, 0.5*cm))
            except:
                pass

        # Titre
        title_style = styles['Heading1']
        title_style.textColor = colors.HexColor('#006b01')
        title_style.alignment = 1  # Centré
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.5*cm))

        # Métadonnées
        if metadata:
            meta_text = " | ".join([f"<b>{k}:</b> {v}" for k, v in metadata.items()])
            elements.append(Paragraph(meta_text, styles['Normal']))
            elements.append(Spacer(1, 0.5*cm))

        # Préparation des données du tableau
        table_data = [[column_labels.get(col, col) for col in columns]]

        for row in data:
            table_data.append([str(row.get(col, '-')) for col in columns])

        # Limitation du nombre de lignes pour le PDF
        if len(table_data) > 101:  # 1 header + 100 lignes
            table_data = table_data[:101]
            elements.append(Paragraph(
                f"<i>Note: Seules les 100 premières lignes sont affichées (total: {len(data)})</i>",
                styles['Normal']
            ))
            elements.append(Spacer(1, 0.3*cm))

        # Création du tableau
        table = Table(table_data)
        table.setStyle(TableStyle([
            # En-tête
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#006b01')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),

            # Corps du tableau
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
        ]))

        elements.append(table)

        # Pied de page
        elements.append(Spacer(1, 1*cm))
        footer_text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} - FPMsigm | Inspections"
        elements.append(Paragraph(footer_text, styles['Italic']))

        # Construction du PDF
        doc.build(elements)

        # Création de la réponse
        buffer.seek(0)
        response = make_response(buffer.read())
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        response.headers["Content-type"] = "application/pdf"

        return response

    @staticmethod
    def prepare_export_data(results, export_format='csv'):
        """
        Prépare les données pour l'export du scénario 1 (modèle corrigé)

        Args:
            results: Résultats de l'analyse
            export_format: Format d'export

        Returns:
            Dict avec données formatées et colonnes
        """
        if not results:
            return {'data': [], 'columns': [], 'column_labels': {}}

        # Colonnes détaillées pour le modèle 1 corrigé (compatible avec admi_claude.py)
        columns = [
            'num_pec',
            'montant_total_pec',
            'LIBELLE_TYPE_PRESTATION',
            'libelle_etat_qualificatif',
            'structure_initiatrice',
            'structure_propose',
            'structure_executante',
            'structure_origine_bulletin',
            'ps_initiateur',
            'tel_initiateur',
            'ps_executant',
            'tel_executant',
            'date_dmd_acte_trans',
            'date_debut_execution',
            'date_fin_execution',
            'date_accuser_reception',
            'cle_validation',
            'nombre_jour_hospitalisation',
            'num_bnf',
            'nom_beneficiaire',
            'prenom_beneficiaire',
            'date_naissance',
            'telephone',
            'sexe'
        ]

        # Labels des colonnes (français)
        column_labels = {
            'num_pec': 'Numéro PEC',
            'montant_total_pec': 'Montant Total (FCFA)',
            'LIBELLE_TYPE_PRESTATION': 'Type Prestation',
            'libelle_etat_qualificatif': 'État Qualificatif',
            'structure_initiatrice': 'Structure Initiatrice',
            'structure_propose': 'Structure Proposée',
            'structure_executante': 'Structure Exécutante',
            'structure_origine_bulletin': 'Structure Origine Bulletin',
            'ps_initiateur': 'Personnel Initiateur',
            'tel_initiateur': 'Tél. Initiateur',
            'ps_executant': 'Personnel Exécutant',
            'tel_executant': 'Tél. Exécutant',
            'date_dmd_acte_trans': 'Date Demande Acte',
            'date_debut_execution': 'Date Début Exécution',
            'date_fin_execution': 'Date Fin Exécution',
            'date_accuser_reception': 'Date Accusé Réception',
            'cle_validation': 'Clé Validation',
            'nombre_jour_hospitalisation': 'Nb Jours Hospitalisation',
            'num_bnf': 'Num. Bénéficiaire',
            'nom_beneficiaire': 'Nom Bénéficiaire',
            'prenom_beneficiaire': 'Prénom Bénéficiaire',
            'date_naissance': 'Date Naissance',
            'telephone': 'Téléphone',
            'sexe': 'Sexe'
        }

        return {
            'data': results,
            'columns': columns,
            'column_labels': column_labels
        }

    @staticmethod
    def prepare_scenario2_export_data(results, filters, export_format='csv'):
        """
        Prépare les données pour l'export du scénario 2

        Args:
            results: Résultats de l'analyse
            filters: Filtres appliqués
            export_format: Format d'export

        Returns:
            Dict avec données formatées et colonnes
        """
        if not results:
            return {'data': [], 'columns': [], 'column_labels': {}}

        # Colonnes à exporter selon les filtres
        columns = ['nom_structure_executante', 'num_pec', 'date_executante_soin']

        if filters.get('show_beneficiaire', True):
            columns.extend(['num_bnf', 'nom_prenom'])

        if filters.get('show_telephone', True):
            columns.append('telephone')

        if filters.get('show_sexe', True):
            columns.append('sexe')

        if filters.get('show_date_naissance', True):
            columns.append('date_naissance')

        if filters.get('show_type_trans', True):
            columns.append('libelle_type_trans')

        if filters.get('show_nb_lignes', True):
            columns.append('nb_lignes')

        columns.extend(['montant_execute_total', 'source_ligne', 'montant_group_numpec'])

        # Labels des colonnes
        column_labels = {
            'id_structure_executante': 'ID Structure',
            'nom_structure_executante': 'Structure exécutante',
            'num_pec': 'Numéro PEC',
            'date_executante_soin': 'Date exécution',
            'num_bnf': 'Num. bénéficiaire',
            'nom_prenom': 'Nom & Prénom',
            'telephone': 'Téléphone',
            'sexe': 'Sexe',
            'date_naissance': 'Date naissance',
            'libelle_type_trans': 'Type transaction',
            'nb_lignes': 'Nb lignes',
            'montant_execute_total': 'Montant total (FCFA)',
            'source_ligne': 'Source(s)',
            'montant_group_numpec': 'Total cumulé PEC (FCFA)'
        }

        # Formatage des données
        formatted_data = []
        for row in results:
            formatted_row = {}
            for col in columns:
                value = row.get(col)
                # Formatage des dates
                if col in ['date_executante_soin', 'date_naissance'] and value:
                    if hasattr(value, 'strftime'):
                        formatted_row[col] = value.strftime('%d/%m/%Y')
                    else:
                        formatted_row[col] = str(value)
                # Formatage des montants
                elif col in ['montant_execute_total', 'montant_group_numpec'] and value:
                    formatted_row[col] = f"{float(value):,.0f}".replace(',', ' ')
                else:
                    formatted_row[col] = value if value is not None else 'N/A'
            formatted_data.append(formatted_row)

        return {
            'data': formatted_data,
            'columns': columns,
            'column_labels': column_labels
        }

    @staticmethod
    def export_to_word(data, columns, column_labels, title, filename_prefix='export',
                      metadata=None, logo_path=None):
        """
        Exporte les données en format Word (DOCX)

        Args:
            data: Liste de dictionnaires
            columns: Liste des clés des colonnes
            column_labels: Dict mapping clés -> libellés
            title: Titre du document
            filename_prefix: Préfixe du nom de fichier
            metadata: Dict avec métadonnées
            logo_path: Chemin vers le logo

        Returns:
            Response Flask avec le fichier DOCX
        """
        timestamp = datetime.now().strftime('%Y-%m-%d_%H%M')
        filename = f"{filename_prefix}_{timestamp}.docx"

        # Création du document
        doc = Document()

        # Logo
        if logo_path and os.path.exists(logo_path):
            try:
                doc.add_picture(logo_path, width=Inches(1.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass

        # Titre
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.color.rgb = RGBColor(0, 107, 1)

        # Métadonnées
        if metadata:
            doc.add_paragraph()
            for key, value in metadata.items():
                para = doc.add_paragraph()
                para.add_run(f"{key}: ").bold = True
                para.add_run(str(value))

        doc.add_paragraph()

        # Limitation du nombre de lignes pour Word
        limited_data = data[:100] if len(data) > 100 else data
        if len(data) > 100:
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(
                f"Note: Seules les 100 premières lignes sont affichées (total: {len(data)})"
            )
            note_run.italic = True

        # Tableau
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Light Grid Accent 1'

        # En-tête du tableau
        header_cells = table.rows[0].cells
        for idx, col_key in enumerate(columns):
            cell = header_cells[idx]
            cell.text = column_labels.get(col_key, col_key)
            # Style de l'en-tête
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Couleur de fond (via shading)
            try:
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn

                tcPr = cell._element.get_or_add_tcPr()
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '006b01')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                tcPr.append(shading_elm)
            except Exception as e:
                # Si le shading échoue, on continue sans couleur de fond
                pass

        # Données du tableau
        for row_data in limited_data:
            row_cells = table.add_row().cells
            for idx, col_key in enumerate(columns):
                value = row_data.get(col_key, '')
                row_cells[idx].text = str(value)
                # Taille de police
                for paragraph in row_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Pied de page
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} - FPMsigm | Inspections"
        )
        footer_run.italic = True
        footer_run.font.size = Pt(9)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sauvegarde en mémoire
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Création de la réponse
        response = make_response(buffer.read())
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return response
