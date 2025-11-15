#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Application Web Streamlit pour la génération d'états synthétiques PEC
=====================================================================

Interface web permettant de :
- Filtrer les PEC par dates et montants
- Visualiser les résultats dans un tableau interactif
- Télécharger les rapports en différents formats (Excel, Word, CSV, PDF)
- Consulter les détails des factures
"""

import streamlit as st
import pandas as pd
import io
from datetime import datetime, date
import sys
import os

# Importer la classe GenerateurEtatPEC depuis admi_claude.py
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from admi_claude import GenerateurEtatPEC

# Imports pour les exports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Configuration de la page
st.set_page_config(
    page_title="États Synthétiques PEC",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS personnalisé
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
    }
    .stButton>button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)


def generer_export_excel(df: pd.DataFrame) -> bytes:
    """
    Génère un fichier Excel formaté avec styles

    Args:
        df: DataFrame à exporter

    Returns:
        Bytes du fichier Excel
    """
    output = io.BytesIO()

    # Créer le fichier Excel avec openpyxl pour un meilleur formatage
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='État Synthétique PEC', index=False)

        # Obtenir le workbook et la feuille
        workbook = writer.book
        worksheet = writer.sheets['État Synthétique PEC']

        # Styles
        header_fill = PatternFill(start_color="1F77B4", end_color="1F77B4", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Formater les en-têtes
        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border

        # Ajuster la largeur des colonnes
        column_widths = {
            'A': 18,  # num_pec
            'B': 18,  # montant_total_pec
            'C': 25,  # LIBELLE_TYPE_PRESTATION
            'D': 25,  # libelle_etat_qualificatif
            'E': 35,  # structure_initiatrice
            'F': 35,  # structure_propose
            'G': 35,  # structure_executante
            'H': 35,  # structure_origine_bulletin
            'I': 25,  # ps_initiateur
            'J': 15,  # tel_initiateur
            'K': 25,  # ps_executant
            'L': 15,  # tel_executant
            'M': 20,  # date_dmd_acte_trans
            'N': 20,  # date_debut_execution
            'O': 20,  # date_fin_execution
            'P': 20,  # date_accuser_reception
            'Q': 15,  # cle_validation
            'R': 15,  # nombre_jour_hospitalisation
            'S': 15,  # num_bnf
            'T': 20,  # nom_beneficiaire
            'U': 20,  # prenom_beneficiaire
            'V': 15,  # date_naissance
            'W': 15,  # telephone
            'X': 8,   # sexe
            'Y': 15,  # Facture
        }

        for col, width in column_widths.items():
            worksheet.column_dimensions[col].width = width

        # Formater les cellules de données
        for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row):
            for cell in row:
                cell.border = border
                cell.alignment = Alignment(vertical='center')

                # Formater les montants
                if cell.column == 2:  # montant_total_pec
                    cell.number_format = '#,##0'
                    cell.alignment = Alignment(horizontal='right', vertical='center')

                # Formater les dates
                elif cell.column in [13, 14, 15, 16, 22]:  # colonnes de dates
                    if cell.value:
                        cell.number_format = 'DD/MM/YYYY HH:MM:SS'

        # Figer la première ligne
        worksheet.freeze_panes = 'A2'

    output.seek(0)
    return output.read()


def generer_export_word(df: pd.DataFrame) -> bytes:
    """
    Génère un fichier Word avec tableau formaté

    Args:
        df: DataFrame à exporter

    Returns:
        Bytes du fichier Word
    """
    doc = Document()

    # Titre
    title = doc.add_heading('État Synthétique des PEC', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Informations du rapport
    doc.add_paragraph(f'Date de génération : {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph(f'Nombre de PEC : {len(df)}')
    doc.add_paragraph(f'Montant total : {df["montant_total_pec"].sum():,.0f} FCFA')
    doc.add_paragraph('')

    # Note : Pour Word, on va créer un tableau simplifié avec les colonnes principales
    # car un tableau avec 25 colonnes sera trop large
    colonnes_principales = [
        'num_pec', 'montant_total_pec', 'LIBELLE_TYPE_PRESTATION',
        'structure_executante', 'date_debut_execution', 'nom_beneficiaire',
        'prenom_beneficiaire', 'Facture'
    ]

    df_simplifie = df[colonnes_principales].copy()

    # Créer le tableau
    table = doc.add_table(rows=1, cols=len(colonnes_principales))
    table.style = 'Light Grid Accent 1'

    # En-têtes
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(colonnes_principales):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Données
    for _, row in df_simplifie.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(colonnes_principales):
            value = row[col]
            if pd.isna(value):
                row_cells[i].text = ''
            elif col == 'montant_total_pec':
                row_cells[i].text = f'{value:,.0f}'
            elif isinstance(value, datetime):
                row_cells[i].text = value.strftime('%d/%m/%Y %H:%M')
            else:
                row_cells[i].text = str(value)

    # Note de bas de page
    doc.add_paragraph('')
    doc.add_paragraph('Note : Tableau simplifié. Téléchargez le fichier Excel pour voir toutes les colonnes.',
                      style='Intense Quote')

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def generer_export_csv(df: pd.DataFrame) -> bytes:
    """
    Génère un fichier CSV

    Args:
        df: DataFrame à exporter

    Returns:
        Bytes du fichier CSV
    """
    output = io.StringIO()
    df.to_csv(output, index=False, encoding='utf-8-sig', sep=';')
    return output.getvalue().encode('utf-8-sig')


def generer_export_pdf(df: pd.DataFrame) -> bytes:
    """
    Génère un fichier PDF avec tableau

    Args:
        df: DataFrame à exporter

    Returns:
        Bytes du fichier PDF
    """
    output = io.BytesIO()

    # Créer le PDF en mode paysage
    doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    elements = []

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1f77b4'),
        alignment=1  # Center
    )

    # Titre
    title = Paragraph("État Synthétique des PEC", title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.3*inch))

    # Informations
    info_style = styles['Normal']
    info_text = f"""
    <b>Date de génération :</b> {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}<br/>
    <b>Nombre de PEC :</b> {len(df)}<br/>
    <b>Montant total :</b> {df['montant_total_pec'].sum():,.0f} FCFA
    """
    elements.append(Paragraph(info_text, info_style))
    elements.append(Spacer(1, 0.3*inch))

    # Tableau simplifié (comme pour Word)
    colonnes_principales = [
        'num_pec', 'montant_total_pec', 'structure_executante',
        'date_debut_execution', 'nom_beneficiaire', 'Facture'
    ]

    df_simplifie = df[colonnes_principales].copy()

    # Préparer les données du tableau
    data = [colonnes_principales]
    for _, row in df_simplifie.iterrows():
        row_data = []
        for col in colonnes_principales:
            value = row[col]
            if pd.isna(value):
                row_data.append('')
            elif col == 'montant_total_pec':
                row_data.append(f'{value:,.0f}')
            elif isinstance(value, datetime):
                row_data.append(value.strftime('%d/%m/%Y'))
            else:
                row_data.append(str(value)[:20])  # Limiter la longueur
        data.append(row_data)

    # Créer le tableau
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
    ]))

    elements.append(table)

    # Note
    elements.append(Spacer(1, 0.3*inch))
    note = Paragraph("<i>Note : Tableau simplifié. Téléchargez le fichier Excel pour voir toutes les colonnes.</i>",
                    styles['Normal'])
    elements.append(note)

    # Générer le PDF
    doc.build(elements)
    output.seek(0)
    return output.read()


def afficher_details_facture(num_pec: str, generateur: GenerateurEtatPEC):
    """
    Affiche les détails d'une facture dans une modale

    Args:
        num_pec: Numéro de PEC
        generateur: Instance du générateur
    """
    st.subheader(f"Détails de la facture : {num_pec}")

    # Récupérer les détails via SQL_TAMPON
    df_details = generateur.executer_sql_tampon(num_pec)

    if df_details.empty:
        st.warning("Aucun détail trouvé pour cette facture")
        return

    # Afficher les informations
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Nombre de lignes", len(df_details))
    with col2:
        total = (df_details['montant'].fillna(0) * df_details['nb'].fillna(1)).sum()
        st.metric("Montant total", f"{total:,.0f} FCFA")
    with col3:
        st.metric("Structures", df_details['nom_structure'].nunique())

    st.write("---")

    # Tableau des détails
    st.dataframe(
        df_details,
        use_container_width=True,
        hide_index=True,
        column_config={
            "montant": st.column_config.NumberColumn(
                "Montant",
                format="%.0f FCFA"
            ),
            "date_execution": st.column_config.DatetimeColumn(
                "Date d'exécution",
                format="DD/MM/YYYY HH:mm"
            )
        }
    )


def main():
    """Fonction principale de l'application Streamlit"""

    # En-tête
    st.markdown('<h1 class="main-header">📊 États Synthétiques PEC</h1>', unsafe_allow_html=True)
    st.markdown("---")

    # Sidebar - Paramètres de filtrage
    with st.sidebar:
        st.header("⚙️ Paramètres de filtrage")

        # Configuration de la base de données
        with st.expander("🔐 Connexion Base de Données", expanded=False):
            db_host = st.text_input("Host", value="192.168.10.214")
            db_port = st.number_input("Port", value=3306, min_value=1, max_value=65535)
            db_name = st.text_input("Database", value="admi")
            db_user = st.text_input("User", value="user_daci")
            db_password = st.text_input("Password", value="Prestige2025", type="password")

        config_db = {
            'host': db_host,
            'database': db_name,
            'user': db_user,
            'password': db_password,
            'port': db_port
        }

        st.markdown("---")

        # Filtres de dates
        st.subheader("📅 Période")
        col1, col2 = st.columns(2)
        with col1:
            date_debut = st.date_input(
                "Date début",
                value=date.today(),
                format="DD/MM/YYYY"
            )
        with col2:
            date_fin = st.date_input(
                "Date fin",
                value=date.today(),
                format="DD/MM/YYYY"
            )

        st.markdown("---")

        # Filtres de montants
        st.subheader("💰 Montants")
        montant_mini = st.number_input(
            "Montant minimum (FCFA)",
            min_value=0,
            value=0,
            step=1000
        )
        montant_maxi_option = st.selectbox(
            "Montant maximum",
            ["Illimité", "Personnalisé"]
        )
        if montant_maxi_option == "Personnalisé":
            montant_maxi = st.number_input(
                "Montant maximum (FCFA)",
                min_value=montant_mini + 1,
                value=1000000,
                step=1000
            )
        else:
            montant_maxi = float('inf')

        st.markdown("---")

        # Bouton de génération
        generer_button = st.button("🚀 Générer l'état", type="primary", use_container_width=True)

    # Initialiser la session state
    if 'df_resultat' not in st.session_state:
        st.session_state.df_resultat = None
    if 'generateur' not in st.session_state:
        st.session_state.generateur = None

    # Génération de l'état
    if generer_button:
        # Validation des dates
        if date_debut > date_fin:
            st.error("❌ La date de début doit être antérieure à la date de fin")
            return

        # Créer le générateur
        with st.spinner("Connexion à la base de données..."):
            generateur = GenerateurEtatPEC(config_db)
            if not generateur.connecter():
                st.error("❌ Impossible de se connecter à la base de données")
                return

        # Générer l'état
        with st.spinner("Génération de l'état synthétique en cours..."):
            df_resultat = generateur.generer_etat_synthetique(
                date_debut.strftime('%Y-%m-%d'),
                date_fin.strftime('%Y-%m-%d'),
                montant_mini,
                montant_maxi
            )

        if df_resultat.empty:
            st.warning("⚠️ Aucun résultat trouvé pour les critères sélectionnés")
            generateur.deconnecter()
            return

        # Ajouter la colonne Facture
        df_resultat['Facture'] = '🔍 Détails'

        # Sauvegarder dans session_state
        st.session_state.df_resultat = df_resultat
        st.session_state.generateur = generateur

        st.success(f"✅ État généré avec succès : {len(df_resultat)} PEC trouvés")

    # Affichage des résultats
    if st.session_state.df_resultat is not None:
        df = st.session_state.df_resultat

        # Métriques
        st.markdown("### 📈 Statistiques")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Nombre de PEC", len(df))
        with col2:
            st.metric("Montant total", f"{df['montant_total_pec'].sum():,.0f} FCFA")
        with col3:
            st.metric("Montant moyen", f"{df['montant_total_pec'].mean():,.0f} FCFA")
        with col4:
            st.metric("Structures", df['structure_executante'].nunique())

        st.markdown("---")

        # Boutons d'export
        st.markdown("### 📥 Télécharger le rapport")
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            excel_data = generer_export_excel(df)
            st.download_button(
                label="📗 Excel (.xlsx)",
                data=excel_data,
                file_name=f"etat_pec_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

        with col2:
            word_data = generer_export_word(df)
            st.download_button(
                label="📘 Word (.docx)",
                data=word_data,
                file_name=f"etat_pec_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        with col3:
            csv_data = generer_export_csv(df)
            st.download_button(
                label="📄 CSV",
                data=csv_data,
                file_name=f"etat_pec_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True
            )

        with col4:
            pdf_data = generer_export_pdf(df)
            st.download_button(
                label="📕 PDF",
                data=pdf_data,
                file_name=f"etat_pec_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        st.markdown("---")

        # Tableau des résultats
        st.markdown("### 📋 Résultats")

        # Configuration des colonnes pour l'affichage
        column_config = {
            "montant_total_pec": st.column_config.NumberColumn(
                "Montant Total",
                format="%.0f FCFA"
            ),
            "date_debut_execution": st.column_config.DatetimeColumn(
                "Date Début Exécution",
                format="DD/MM/YYYY HH:mm"
            ),
            "date_fin_execution": st.column_config.DatetimeColumn(
                "Date Fin Exécution",
                format="DD/MM/YYYY HH:mm"
            ),
            "date_dmd_acte_trans": st.column_config.DatetimeColumn(
                "Date Demande",
                format="DD/MM/YYYY HH:mm"
            ),
            "date_accuser_reception": st.column_config.DatetimeColumn(
                "Date Accusé Réception",
                format="DD/MM/YYYY HH:mm"
            ),
            "date_naissance": st.column_config.DatetimeColumn(
                "Date Naissance",
                format="DD/MM/YYYY"
            ),
            "Facture": st.column_config.TextColumn(
                "Facture",
                width="small"
            )
        }

        # Affichage du tableau avec possibilité de sélection
        event = st.dataframe(
            df,
            use_container_width=True,
            hide_index=True,
            column_config=column_config,
            selection_mode="single-row",
            on_select="rerun",
            key="data_table"
        )

        # Si une ligne est sélectionnée, afficher les détails
        if len(event.selection.rows) > 0:
            selected_idx = event.selection.rows[0]
            num_pec_selected = df.iloc[selected_idx]['num_pec']

            st.markdown("---")
            with st.expander(f"🔍 Détails de la facture {num_pec_selected}", expanded=True):
                afficher_details_facture(num_pec_selected, st.session_state.generateur)

    else:
        # Message d'accueil
        st.info("👈 Configurez les paramètres dans la barre latérale et cliquez sur 'Générer l'état'")

        st.markdown("""
        ### 📖 Guide d'utilisation

        1. **Connexion** : Vérifiez les paramètres de connexion à la base de données
        2. **Période** : Sélectionnez la plage de dates pour filtrer les PEC
        3. **Montants** : Définissez les limites de montants (optionnel)
        4. **Génération** : Cliquez sur le bouton "Générer l'état"
        5. **Export** : Téléchargez les résultats dans le format de votre choix
        6. **Détails** : Cliquez sur une ligne pour voir les détails de la facture

        ### 📊 Formats d'export disponibles

        - **Excel** : Fichier complet avec toutes les colonnes et formatage
        - **Word** : Document avec tableau simplifié (colonnes principales)
        - **CSV** : Fichier texte pour import dans d'autres outils
        - **PDF** : Document imprimable avec tableau simplifié
        """)


if __name__ == "__main__":
    main()
